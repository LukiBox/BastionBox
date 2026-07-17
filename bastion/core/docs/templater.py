"""Fill a business .docx template — keep the branding, replace the placeholders.

The user's company template carries the logo, fonts, headers/footers, and table
styling. The agent must never rebuild that from scratch; it *fills* it:

* ``{{KEY}}``          → replaced with text (multi-line values become real
                          line breaks, not literal ``\\n``)
* ``{{IMG:name}}``     → replaced with a picture (test photos), sized to fit
* ``{{TABLE:name}}``   → the table row containing it becomes a prototype; one
                          styled row is emitted per data row (test data)

Placeholders are found across *runs* (Word happily splits ``{{TITLE}}`` into
three runs mid-word), in body paragraphs, tables, headers, and footers. The
first run of a matched paragraph keeps its formatting — that's what preserves
the template's look. Unmatched placeholders are left visibly in place and
reported, never silently dropped: a half-filled report must say so.
"""
from __future__ import annotations

import copy
import io
import re
from pathlib import Path
from typing import Any

from .extract import ExtractionUnavailable

_PLACEHOLDER = re.compile(r"\{\{\s*([A-Za-z0-9 _.:-]+?)\s*\}\}")
_IMG_PREFIX = "IMG:"
_TABLE_PREFIX = "TABLE:"


def find_placeholders(template_path: str | Path) -> list[str]:
    """List every ``{{...}}`` key in the template (body, tables, headers, footers)."""
    doc = _open(template_path)
    found: list[str] = []
    for par in _all_paragraphs(doc):
        for m in _PLACEHOLDER.finditer(_par_text(par)):
            key = m.group(1).strip()
            if key not in found:
                found.append(key)
    return found


def fill_docx_template(template_path: str | Path,
                       fields: dict[str, Any] | None = None,
                       images: dict[str, str | Path] | None = None,
                       tables: dict[str, list[list[Any]]] | None = None,
                       ) -> tuple[bytes, list[str]]:
    """Render the template with *fields*/*images*/*tables* filled in.

    Returns ``(docx_bytes, leftover)`` where *leftover* lists placeholders that
    had no value supplied — surfaced to the model so it can fix its call.
    """
    fields = {str(k).strip(): v for k, v in (fields or {}).items()}
    images = {str(k).strip(): v for k, v in (images or {}).items()}
    tables = {str(k).strip(): v for k, v in (tables or {}).items()}
    doc = _open(template_path)

    for table in _all_tables(doc):
        _fill_table_rows(table, tables)
    for par in _all_paragraphs(doc):
        _fill_paragraph(par, fields, images)

    leftover: list[str] = []
    for par in _all_paragraphs(doc):
        for m in _PLACEHOLDER.finditer(_par_text(par)):
            key = m.group(1).strip()
            if key not in leftover:
                leftover.append(key)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), leftover


# ---------------------------------------------------------------------------
# internals
# ---------------------------------------------------------------------------
def _open(template_path: str | Path):
    try:
        import docx
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Filling .docx templates needs python-docx: "
            "pip install python-docx") from exc
    return docx.Document(str(template_path))


def _all_paragraphs(doc):
    """Every paragraph in the document: body, tables, headers, footers."""
    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from walk_tables(cell.tables)

    yield from doc.paragraphs
    yield from walk_tables(doc.tables)
    for section in doc.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            yield from walk_tables(part.tables)


def _all_tables(doc):
    yield from doc.tables
    for section in doc.sections:
        yield from section.header.tables
        yield from section.footer.tables


def _par_text(par) -> str:
    return "".join(r.text for r in par.runs) or par.text


def _fill_paragraph(par, fields: dict[str, Any],
                    images: dict[str, str | Path]) -> None:
    text = _par_text(par)
    if "{{" not in text:
        return

    # A paragraph that *is* an image placeholder becomes the picture itself.
    m = _PLACEHOLDER.fullmatch(text.strip())
    if m and m.group(1).strip().startswith(_IMG_PREFIX):
        name = m.group(1).strip()[len(_IMG_PREFIX):].strip()
        if name in images:
            _replace_with_picture(par, images[name])
        return

    def sub(match: re.Match) -> str:
        key = match.group(1).strip()
        if key.startswith(_IMG_PREFIX) or key.startswith(_TABLE_PREFIX):
            return match.group(0)          # handled elsewhere / leftover
        if key in fields:
            return "" if fields[key] is None else str(fields[key])
        return match.group(0)              # unmatched: keep visible

    new = _PLACEHOLDER.sub(sub, text)
    if new == text:
        return
    # Collapse into the first run (keeps its character formatting — the
    # template's font/size/color for that spot), clear the rest.
    if not par.runs:
        par.text = new
        return
    first, rest = par.runs[0], par.runs[1:]
    _set_run_multiline(first, new)
    for r in rest:
        r.text = ""


def _set_run_multiline(run, value: str) -> None:
    """Set run text where ``\\n`` becomes a real line break in Word."""
    lines = str(value).split("\n")
    run.text = lines[0]
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _replace_with_picture(par, image_path: str | Path) -> None:
    from docx.shared import Inches
    for r in par.runs:
        r.text = ""
    run = par.runs[0] if par.runs else par.add_run()
    run.add_picture(str(image_path), width=Inches(5.5))


def _fill_table_rows(table, tables: dict[str, list[list[Any]]]) -> None:
    """Expand ``{{TABLE:name}}`` prototype rows into one styled row per record."""
    for row in list(table.rows):
        proto_key = None
        for cell in row.cells:
            m = _PLACEHOLDER.search(cell.text)
            if m and m.group(1).strip().startswith(_TABLE_PREFIX):
                proto_key = m.group(1).strip()[len(_TABLE_PREFIX):].strip()
                break
        if proto_key is None or proto_key not in tables:
            continue
        data = tables[proto_key]
        anchor = row._tr
        for record in data:
            clone = copy.deepcopy(anchor)
            anchor.addprevious(clone)
            from docx.table import _Row
            new_row = _Row(clone, table)
            for i, cell in enumerate(new_row.cells):
                value = "" if i >= len(record) or record[i] is None else str(record[i])
                # Keep the prototype cell's paragraph style; swap the text.
                first_par = cell.paragraphs[0]
                if first_par.runs:
                    _set_run_multiline(first_par.runs[0], value)
                    for r in first_par.runs[1:]:
                        r.text = ""
                else:
                    first_par.text = value
                for extra in cell.paragraphs[1:]:
                    for r in extra.runs:
                        r.text = ""
        anchor.getparent().remove(anchor)   # drop the prototype row
