"""Local document extraction — read PDFs, Word, Excel, and text, offline.

This is what lets BastionBox *read a 400-page datasheet* and turn it into a
written answer without a byte leaving the machine. Extraction is entirely local:
PyMuPDF (or pypdf as a fallback) for PDF, python-docx for Word, openpyxl for
Excel, and a plain read for text/markdown/code.

Every backend is an optional dependency imported lazily, so importing this module
never fails; a missing library yields a clear, actionable message instead of a
traceback. PDFs return per-page text so the agent can cite and target pages of a
long datasheet.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

# Extensions we know how to turn into text.
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
XLSX_SUFFIXES = {".xlsx", ".xlsm"}
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".py",
                 ".js", ".ts", ".c", ".h", ".cpp", ".java", ".go", ".rs",
                 ".sql", ".yaml", ".yml", ".toml", ".ini", ".log"}


class ExtractionUnavailable(RuntimeError):
    """Raised when the library needed for a document type isn't installed."""


@dataclass
class DocumentText:
    path: str
    kind: str                 # "pdf" | "docx" | "xlsx" | "text"
    text: str
    pages: list[str] = field(default_factory=list)   # per-page, for PDFs
    page_count: int = 0
    meta: dict = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return len(self.text)


def supported(path: str | Path) -> bool:
    s = Path(path).suffix.lower()
    return (s in PDF_SUFFIXES or s in DOCX_SUFFIXES or s in XLSX_SUFFIXES
            or s in TEXT_SUFFIXES)


def extract(path: str | Path) -> DocumentText:
    """Extract text from *path* by type. Raises ExtractionUnavailable on missing deps."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in PDF_SUFFIXES:
        return _extract_pdf(p)
    if suffix in DOCX_SUFFIXES:
        return _extract_docx(p)
    if suffix in XLSX_SUFFIXES:
        return _extract_xlsx(p)
    return _extract_text(p)


def _extract_pdf(p: Path) -> DocumentText:
    # Prefer PyMuPDF (fast, robust on datasheets); fall back to pypdf.
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(p))
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return DocumentText(str(p), "pdf", "\n".join(pages), pages, len(pages),
                            {"engine": "pymupdf"})
    except ImportError:
        pass
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(p))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return DocumentText(str(p), "pdf", "\n".join(pages), pages, len(pages),
                            {"engine": "pypdf"})
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Reading PDFs needs PyMuPDF or pypdf. Install one: "
            "pip install pymupdf") from exc


def _extract_docx(p: Path) -> DocumentText:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Reading .docx needs python-docx: pip install python-docx") from exc
    d = docx.Document(str(p))
    parts = [para.text for para in d.paragraphs]
    for t, table in enumerate(d.tables, 1):
        parts.append(f"[table {t}]")
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    return DocumentText(str(p), "docx", text, meta={"tables": len(d.tables)})


def _extract_xlsx(p: Path) -> DocumentText:
    try:
        import openpyxl
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Reading .xlsx needs openpyxl: pip install openpyxl") from exc
    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                parts.append("\t".join(cells))
    wb.close()
    return DocumentText(str(p), "xlsx", "\n".join(parts),
                        meta={"sheets": len(wb.sheetnames)})


def _extract_text(p: Path) -> DocumentText:
    raw = p.read_bytes()[:2_000_000]
    return DocumentText(str(p), "text", raw.decode("utf-8", errors="replace"))
