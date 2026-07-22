"""Local document *writing* — the AI produces real .docx / .xlsx / .pdf files.

The other half of the office story: after BastionBox reads a long datasheet
(see :mod:`bastion.core.docs.extract`), it writes structured output the user can
open in Word, Excel, or a PDF viewer — entirely offline. Each writer takes a
simple, model-friendly structure (a list of blocks / rows) so a small local model
can drive it reliably, and returns the bytes to write; the *tool layer* is what
actually commits them through the path jail with a preview.

All heavy libraries are optional and imported lazily, so importing this module is
free and a missing dependency yields a clear message, never a traceback.
"""
from __future__ import annotations

import io
import re
from typing import Any, Sequence
from xml.sax.saxutils import escape as _xml_escape

from .extract import ExtractionUnavailable

_EMPHASIS_BOLD = re.compile(r"\*\*(.+?)\*\*")
_EMPHASIS_ITALIC = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")


def _pad_rows(rows: Sequence[Sequence[Any]]) -> list[list[Any]]:
    """Normalize a possibly-jagged table: every row padded to the widest row.

    Small local models routinely emit tables whose rows differ in cell count;
    without padding the docx writer raised IndexError and the xlsx column
    widths misaligned.
    """
    rows = [list(r) for r in rows]
    width = max((len(r) for r in rows), default=0)
    return [r + [""] * (width - len(r)) for r in rows]


def _image_fit_inches(data: bytes, max_inches: float = 5.8) -> float:
    """Width (inches) to render *data* at: natural size, capped at *max_inches*.

    Blowing a small photo up to full page width looks broken; keeping natural
    size for small images and capping large ones is what a human would do.
    """
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(io.BytesIO(data)) as im:
            w_px = im.size[0]
            dpi = (im.info.get("dpi") or (96, 96))[0] or 96
        return min(max_inches, w_px / float(dpi))
    except Exception:  # noqa: BLE001 - sizing is cosmetic; the cap is safe
        return max_inches


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------
def build_docx(blocks: Sequence[dict[str, Any]]) -> bytes:
    """Render *blocks* to .docx bytes.

    Each block is ``{"type": ..., "text"/"rows": ...}`` where type is one of:
    ``heading`` (+``level`` 1-3), ``paragraph``, ``bullet``, ``table``
    (``rows``: list of lists). This tiny schema is easy for a local model to emit
    and covers the reports/summaries the office flow needs.
    """
    try:
        import docx
        from docx.shared import Pt
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Writing .docx needs python-docx: pip install python-docx") from exc
    def _add_runs(paragraph, text: str) -> None:
        """Add text as runs, honoring **bold** / *italic* markdown emphasis."""
        pos = 0
        for m in re.finditer(r"\*\*(.+?)\*\*|(?<!\*)\*([^*\n]+)\*(?!\*)", text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            if m.group(1) is not None:
                paragraph.add_run(m.group(1)).bold = True
            else:
                paragraph.add_run(m.group(2)).italic = True
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    doc = docx.Document()
    for block in blocks:
        kind = block.get("type", "paragraph")
        if kind == "heading":
            level = max(1, min(9, int(block.get("level", 1) or 1)))
            text = _EMPHASIS_BOLD.sub(r"\1", block.get("text", ""))
            doc.add_heading(text, level=level)
        elif kind == "bullet":
            _add_runs(doc.add_paragraph(style="List Bullet"), block.get("text", ""))
        elif kind == "table":
            rows = _pad_rows(block.get("rows", []))
            if rows and rows[0]:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, val in enumerate(row):
                        table.rows[r].cells[c].text = str(val)
        elif kind == "chart":
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
            from . import charts as _charts
            spec = block.get("spec") or {}
            try:
                png = _charts.chart_png(spec)
                doc.add_picture(io.BytesIO(png), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:  # noqa: BLE001 - no raster engine or bad spec:
                # fall back to the chart's data as a table so the doc still saves.
                try:
                    rows = _charts.chart_table(spec)
                except Exception:  # noqa: BLE001 - spec unusable; skip the block
                    rows = []
                if rows and rows[0]:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Light Grid Accent 1"
                    for r, row in enumerate(rows):
                        for c, val in enumerate(row):
                            table.rows[r].cells[c].text = str(val)
        elif kind == "image":
            data = block.get("data")
            if data:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.shared import Inches
                doc.add_picture(io.BytesIO(data),
                                width=Inches(_image_fit_inches(data)))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if block.get("caption"):
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap.add_run(str(block["caption"]))
                    run.italic = True
                    run.font.size = Pt(9)
        else:  # paragraph
            _add_runs(doc.add_paragraph(), block.get("text", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Excel (.xlsx)
# ---------------------------------------------------------------------------
def build_xlsx(sheets: Sequence[dict[str, Any]]) -> bytes:
    """Render *sheets* to .xlsx bytes.

    Each sheet is ``{"name": str, "rows": list[list]}``; the first row is styled
    as a bold header. A single flat ``{"rows": [...]}`` is also accepted for the
    common one-table case. An optional ``"chart": {"type": "bar"|"line"|"pie",
    "title": str}`` adds a native Excel chart built from the sheet's own data
    (first column = categories, remaining columns = series) — it stays live and
    editable in Excel.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Writing .xlsx needs openpyxl: pip install openpyxl") from exc
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for i, sheet in enumerate(sheets):
        ws = wb.create_sheet(_sheet_title(sheet.get("name"), i))
        rows = _pad_rows(sheet.get("rows", []))
        for r, row in enumerate(rows, 1):
            for c, val in enumerate(row, 1):
                # Coerce numeric-looking strings (data rows only) so Excel
                # gets real numbers it can sum/chart, not text.
                cell = ws.cell(row=r, column=c,
                               value=_coerce(val) if r > 1 else val)
                if r == 1:  # header styling
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor="54806C")
        ws.freeze_panes = "A2"  # keep the header visible while scrolling
        # Auto-ish column widths from the longest cell.
        for c, col in enumerate(zip(*rows), 1) if rows else []:
            width = min(60, max((len(str(v)) for v in col), default=10) + 2)
            ws.column_dimensions[openpyxl.utils.get_column_letter(c)].width = width
        if sheet.get("chart") and len(rows) >= 2 and len(rows[0]) >= 2:
            _add_xlsx_chart(ws, rows, sheet["chart"])
    if not wb.sheetnames:
        wb.create_sheet("Sheet1")
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _add_xlsx_chart(ws, rows: list[list[Any]], spec: Any) -> None:
    """Add a native, still-editable Excel chart built from the sheet's data.

    Column A (below the header) supplies the categories; every remaining
    column becomes a series (pie uses only the first). Best-effort: a chart
    that cannot be built must never cost the user their data sheet.
    """
    try:
        from openpyxl.chart import BarChart, LineChart, PieChart, Reference
        from openpyxl.utils import get_column_letter

        kind = str((spec or {}).get("type", "bar")).strip().lower() \
            if isinstance(spec, dict) else str(spec).strip().lower()
        title = (spec or {}).get("title", "") if isinstance(spec, dict) else ""
        n_rows, n_cols = len(rows), len(rows[0])
        if kind == "pie":
            chart = PieChart()
            data = Reference(ws, min_col=2, max_col=2, min_row=1, max_row=n_rows)
        else:
            chart = LineChart() if kind == "line" else BarChart()
            data = Reference(ws, min_col=2, max_col=n_cols,
                             min_row=1, max_row=n_rows)
        cats = Reference(ws, min_col=1, min_row=2, max_row=n_rows)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(cats)
        chart.title = str(title) or None
        chart.style = 10
        chart.height, chart.width = 8, 14   # cm
        ws.add_chart(chart, f"{get_column_letter(n_cols + 2)}2")
    except Exception:  # noqa: BLE001 - the data always outranks the chart
        pass


def _sheet_title(name: Any, index: int) -> str:
    """A sheet title Excel accepts: strip the chars openpyxl rejects
    (``[]:*?/\\``), collapse to 31 chars, never empty."""
    title = re.sub(r"[\[\]:*?/\\]", " ", str(name or "")).strip()[:31]
    return title or f"Sheet{index + 1}"


def _coerce(val: Any):
    """Turn a numeric-looking string into int/float; leave everything else.

    Deliberately conservative: leading zeros ("007"), IDs with spaces, and
    anything ambiguous stay strings.
    """
    if not isinstance(val, str):
        return val
    s = val.strip()
    if not s or re.match(r"^-?0\d", s):  # "007" is an ID, not the number 7
        return val
    try:
        return int(s)
    except ValueError:
        pass
    try:
        f = float(s)
        return f if s.lower() not in ("nan", "inf", "-inf", "infinity", "-infinity") else val
    except ValueError:
        return val


# ---------------------------------------------------------------------------
# PDF (.pdf)
# ---------------------------------------------------------------------------
def build_pdf(blocks: Sequence[dict[str, Any]], title: str = "") -> bytes:
    """Render *blocks* (same schema as :func:`build_docx`) to PDF bytes."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, ListFlowable, ListItem)
        from reportlab.lib import colors
        from reportlab.lib.units import inch
    except ImportError as exc:
        raise ExtractionUnavailable(
            "Writing PDFs needs reportlab: pip install reportlab") from exc
    def _rl_text(text: str) -> str:
        """Escape for reportlab's XML-ish Paragraph markup, then re-inject
        <b>/<i> from markdown emphasis. Raw '<' or '&' in model output used to
        crash the PDF build with a parse error."""
        safe = _xml_escape(text)
        safe = _EMPHASIS_BOLD.sub(r"<b>\1</b>", safe)
        return _EMPHASIS_ITALIC.sub(r"<i>\1</i>", safe)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=title or "BastionBox")
    styles = getSampleStyleSheet()
    flow: list = []
    if title:
        flow.append(Paragraph(_rl_text(title), styles["Title"]))
        flow.append(Spacer(1, 0.2 * inch))
    for block in blocks:
        kind = block.get("type", "paragraph")
        if kind == "heading":
            lvl = max(1, min(3, int(block.get("level", 1) or 1)))
            flow.append(Paragraph(_rl_text(block.get("text", "")),
                                  styles[f"Heading{lvl}"]))
        elif kind == "bullet":
            flow.append(ListFlowable(
                [ListItem(Paragraph(_rl_text(block.get("text", "")),
                                    styles["BodyText"]))],
                bulletType="bullet"))
        elif kind == "table":
            rows = [[str(c) for c in r] for r in _pad_rows(block.get("rows", []))]
            if rows:
                t = Table(rows, repeatRows=1)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#54806C")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                ]))
                flow.append(t)
        elif kind == "chart":
            from . import charts as _charts
            spec = block.get("spec") or {}
            try:
                drawing = _charts.build_drawing(spec)
                drawing.hAlign = "CENTER"     # true vector art in the PDF
                flow.append(drawing)
            except Exception:  # noqa: BLE001 - bad spec: degrade to a data table
                try:
                    rows = _charts.chart_table(spec)
                except Exception:  # noqa: BLE001 - spec unusable; skip the block
                    rows = []
                if rows and rows[0]:
                    t = Table(rows, hAlign="CENTER")
                    t.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#54806C")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ]))
                    flow.append(t)
        elif kind == "image":
            data = block.get("data")
            if data:
                from reportlab.platypus import Image as RLImage
                img = RLImage(io.BytesIO(data))
                max_w = _image_fit_inches(data, 5.9) * inch
                if img.imageWidth > max_w:
                    img.drawHeight = img.imageHeight * max_w / img.imageWidth
                    img.drawWidth = max_w
                img.hAlign = "CENTER"
                flow.append(img)
                if block.get("caption"):
                    from reportlab.lib.styles import ParagraphStyle
                    cap_style = ParagraphStyle(
                        "BastionCaption", parent=styles["BodyText"],
                        alignment=1, fontSize=9)   # 1 = TA_CENTER
                    flow.append(Paragraph(
                        f"<i>{_rl_text(str(block['caption']))}</i>", cap_style))
        else:
            flow.append(Paragraph(_rl_text(block.get("text", "")),
                                  styles["BodyText"]))
        flow.append(Spacer(1, 0.08 * inch))
    doc.build(flow)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HTML (self-contained report — charts/photos embedded as base64)
# ---------------------------------------------------------------------------
def build_html(blocks: Sequence[dict[str, Any]], title: str = "") -> bytes:
    """Render *blocks* to a single self-contained .html report.

    Same block schema as build_docx/build_pdf. Charts render through the same
    tested pipeline (reportlab → PNG) and land as base64 ``data:`` images, as do
    photos — the file has zero external references, so it opens anywhere and can
    be mailed to an investor as-is. All text is HTML-escaped; only the **bold**
    and *italic* emphasis markers become markup.
    """
    import base64
    import html as _html

    def _fmt(text: str) -> str:
        out = _html.escape(str(text), quote=False)
        out = _EMPHASIS_BOLD.sub(r"<strong>\1</strong>", out)
        out = _EMPHASIS_ITALIC.sub(r"<em>\1</em>", out)
        return out

    def _img_mime(data: bytes) -> str:
        if data[:8] == b"\x89PNG\r\n\x1a\n":
            return "image/png"
        if data[:2] == b"\xff\xd8":
            return "image/jpeg"
        if data[:6] in (b"GIF87a", b"GIF89a"):
            return "image/gif"
        if data[:2] == b"BM":
            return "image/bmp"
        return "image/png"

    parts: list[str] = []
    open_list = False

    def _close_list() -> None:
        nonlocal open_list
        if open_list:
            parts.append("</ul>")
            open_list = False

    for block in blocks:
        kind = block.get("type", "paragraph")
        if kind == "bullet":
            if not open_list:
                parts.append("<ul>")
                open_list = True
            parts.append(f"<li>{_fmt(block.get('text', ''))}</li>")
            continue
        _close_list()
        if kind == "heading":
            level = max(1, min(4, int(block.get("level", 1) or 1)))
            parts.append(f"<h{level}>{_fmt(block.get('text', ''))}</h{level}>")
        elif kind == "table":
            rows = _pad_rows(block.get("rows", []))
            if rows and rows[0]:
                head = "".join(f"<th>{_fmt(v)}</th>" for v in rows[0])
                body = "".join(
                    "<tr>" + "".join(f"<td>{_fmt(v)}</td>" for v in row) + "</tr>"
                    for row in rows[1:])
                parts.append(f"<table><thead><tr>{head}</tr></thead>"
                             f"<tbody>{body}</tbody></table>")
        elif kind == "chart":
            from . import charts as _charts
            spec = block.get("spec") or {}
            try:
                png = _charts.chart_png(spec)
                b64 = base64.b64encode(png).decode("ascii")
                parts.append(f'<figure><img alt="chart" '
                             f'src="data:image/png;base64,{b64}"></figure>')
            except Exception:  # noqa: BLE001 - no render engine or bad spec:
                # degrade the chart to its data table so the report still writes
                # (the numbers always outrank the picture).
                try:
                    rows = _charts.chart_table(spec)
                except Exception:  # noqa: BLE001 - spec unusable; skip the block
                    continue
                head = "".join(f"<th>{_fmt(v)}</th>" for v in rows[0])
                body = "".join(
                    "<tr>" + "".join(f"<td>{_fmt(v)}</td>" for v in row) + "</tr>"
                    for row in rows[1:])
                parts.append(f"<table><thead><tr>{head}</tr></thead>"
                             f"<tbody>{body}</tbody></table>")
        elif kind == "image":
            data = block.get("data")
            if data:
                b64 = base64.b64encode(data).decode("ascii")
                cap = (f"<figcaption>{_fmt(block['caption'])}</figcaption>"
                       if block.get("caption") else "")
                parts.append(f'<figure><img alt="" src="data:{_img_mime(data)};'
                             f'base64,{b64}">{cap}</figure>')
        else:
            parts.append(f"<p>{_fmt(block.get('text', ''))}</p>")
    _close_list()

    css = """
    body { font-family: 'Segoe UI', system-ui, sans-serif; color: #24312b;
           max-width: 860px; margin: 2.2rem auto; padding: 0 1.2rem;
           line-height: 1.55; background: #fbfcfb; }
    h1, h2, h3, h4 { color: #375447; line-height: 1.25; }
    h1 { border-bottom: 3px solid #6FA28E; padding-bottom: .35rem; }
    table { border-collapse: collapse; margin: 1rem 0; width: 100%; }
    th { background: #54806C; color: #fff; text-align: left; }
    th, td { border: 1px solid #cfdcd6; padding: .45rem .7rem; }
    tr:nth-child(even) td { background: #f1f6f3; }
    figure { text-align: center; margin: 1.4rem 0; }
    figure img { max-width: 100%; height: auto; }
    figcaption { font-style: italic; font-size: .9rem; color: #5c6f66; }
    @media print { body { margin: 0; background: #fff; } }
    """.strip()
    doc_title = _fmt(title) if title else "Report"
    head_h1 = f"<h1>{_fmt(title)}</h1>" if title else ""
    page = (f"<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
            f"<title>{doc_title}</title><style>{css}</style></head>"
            f"<body>{head_h1}{''.join(parts)}</body></html>")
    return page.encode("utf-8")
